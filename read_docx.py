#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import os
from docx import Document

# Set encoding for stdout
sys.stdout.reconfigure(encoding='utf-8')

# Find the docx file
for file in os.listdir('.'):
    if file.endswith('.docx'):
        docx_file = file
        break

print(f"Reading file: {docx_file}")

try:
    doc = Document(docx_file)
    full_text = []

    # Read paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Read tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())

    # Print all text
    for text in full_text:
        print(text)
        print("-" * 50)

except Exception as e:
    print(f"Error: {e}")
